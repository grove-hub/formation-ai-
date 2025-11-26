import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
import os 
import re 
from io import BytesIO
from pypdf import PdfReader 
from docx import Document

try:
    from azure.identity import ClientSecretCredential, DefaultAzureCredential
    from azure.storage.filedatalake import DataLakeServiceClient
except ModuleNotFoundError:
    raise SystemExit(
        "SDK Azure manquant: installez 'azure-identity' et 'azure-storage-file-datalake'.\n"
        "Exemples:\n"
        "  - pip:    python -m pip install azure-identity azure-storage-file-datalake\n"
        "  - conda:  conda install -c conda-forge azure-identity azure-storage-file-datalake"
    )

# Recupere les document des url passe et les transforme en texte netoyer et pret pour le pipeline

# ---------- CONFIGURATION ADLS ----------
# Variables d'environnement attendues:
# - AZURE_STORAGE_ACCOUNT: nom du compte (sans suffixe .dfs.core.windows.net)
# - AZURE_STORAGE_KEY: clé de compte (option 1 d'authentification)
# - STORAGE_FILESYSTEM: nom du file system (ex: 'data')
# - Option service principal (si pas de STORAGE_KEY):
#   AZURE_TENANT_ID, AZURE_CLIENT_ID, AZURE_CLIENT_SECRET
ACCOUNT_NAME = "juridicai"
ACCOUNT_KEY = os.getenv("AZURE_STORAGE_KEY", "").strip()
FILESYSTEM = "data"

RAW_DIR = "raw_pdfs"
BEFORE_CLEAN_DIR = "before_clean_data"
CLEAN_DIR = "clean_data"
# ---------------------------------------

def get_dls_client():
    """Crée et retourne un client Azure Data Lake Storage"""
    if not ACCOUNT_NAME or not FILESYSTEM:
        raise SystemExit("Veuillez définir AZURE_STORAGE_ACCOUNT et STORAGE_FILESYSTEM.")
    account_url = f"https://{ACCOUNT_NAME}.dfs.core.windows.net"

    if ACCOUNT_KEY:
        return DataLakeServiceClient(account_url=account_url, credential=ACCOUNT_KEY)

    # Essayer DefaultAzureCredential (Managed Identity / dev env), sinon service principal
    try:
        credential = DefaultAzureCredential(exclude_interactive_browser_credential=False)
        return DataLakeServiceClient(account_url=account_url, credential=credential)
    except Exception:
        tenant_id = os.getenv("AZURE_TENANT_ID")
        client_id = os.getenv("AZURE_CLIENT_ID")
        client_secret = os.getenv("AZURE_CLIENT_SECRET")
        if not (tenant_id and client_id and client_secret):
            raise SystemExit(
                "Aucun mode d'authentification disponible. Fournissez AZURE_STORAGE_KEY "
                "ou un service principal (AZURE_TENANT_ID, AZURE_CLIENT_ID, AZURE_CLIENT_SECRET)."
            )
        credential = ClientSecretCredential(tenant_id=tenant_id, client_id=client_id, client_secret=client_secret)
        return DataLakeServiceClient(account_url=account_url, credential=credential)

def list_files_in_adls(file_system_client, directory_path):
    """Liste tous les fichiers dans un répertoire ADLS"""
    try:
        paths_iter = file_system_client.get_paths(path=directory_path)
        files = []
        for p in paths_iter:
            if not p.is_directory and p.name.startswith(directory_path + "/"):
                files.append(os.path.basename(p.name))
        return files
    except Exception:
        return []

def upload_file_to_adls(file_system_client, file_path, file_data):
    """Upload un fichier vers ADLS"""
    try:
        file_client = file_system_client.get_file_client(file_path)
        # Supprimer si existe déjà
        try:
            file_client.delete_file()
        except Exception:
            pass
        # Créer et écrire
        file_client.create_file()
        file_client.append_data(data=file_data, offset=0, length=len(file_data))
        file_client.flush_data(len(file_data))
        return True
    except Exception as e:
        print(f"Erreur lors de l'upload: {e}")
        return False

def download_file_from_adls(file_system_client, file_path):
    """Télécharge un fichier depuis ADLS et retourne les bytes"""
    try:
        file_client = file_system_client.get_file_client(file_path)
        if hasattr(file_client, "read_file"):
            downloader = file_client.read_file()
        else:
            downloader = file_client.download_file()
        return downloader.readall()
    except Exception as e:
        print(f"Erreur lors du téléchargement: {e}")
        return None

def read_text_from_adls(file_system_client, file_path):
    """Lit un fichier texte depuis ADLS et retourne son contenu"""
    try:
        file_client = file_system_client.get_file_client(file_path)
        if hasattr(file_client, "read_file"):
            downloader = file_client.read_file()
        else:
            downloader = file_client.download_file()
        return downloader.readall().decode("utf-8")
    except Exception as e:
        print(f"Erreur lors de la lecture: {e}")
        return None

def write_text_to_adls(file_system_client, file_path, text):
    """Écrit un texte dans un fichier ADLS"""
    try:
        file_client = file_system_client.get_file_client(file_path)
        # Supprimer si existe déjà
        try:
            file_client.delete_file()
        except Exception:
            pass
        # Créer et écrire
        data = text.encode("utf-8")
        file_client.create_file()
        file_client.append_data(data=data, offset=0, length=len(data))
        file_client.flush_data(len(data))
        return True
    except Exception as e:
        print(f"Erreur lors de l'écriture: {e}")
        return False

def delete_file_from_adls(file_system_client, file_path):
    """Supprime un fichier depuis ADLS"""
    try:
        file_client = file_system_client.get_file_client(file_path)
        file_client.delete_file()
        return True
    except Exception as e:
        print(f"Erreur lors de la suppression: {e}")
        return False

class TextScrapper():
    def __init__(self, url):
        
        try:
            self.headers = {
                # Utilisez un User-Agent commun pour Chrome, Firefox, ou autre
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.127 Safari/537.36'
            }
            self.url = url
            self.response = requests.get(self.url, self.headers)
            self.response.raise_for_status()

            self.html = self.response.text

            self.soup = BeautifulSoup(self.html, "html.parser")
        except:
            self.soup = ""

        self.pdf_urls = []    

        # Initialiser le client Azure Data Lake Storage
        self.dls_client = get_dls_client()
        self.file_system = self.dls_client.get_file_system_client(FILESYSTEM)
        
        # Chemins Azure pour les différents dossiers
        self.raw_pdf = RAW_DIR
        self.output_folder = BEFORE_CLEAN_DIR
        self.final_folder = CLEAN_DIR
        
        # Créer les dossiers dans Azure si ils n'existent pas
        for directory in [self.raw_pdf, self.output_folder, self.final_folder]:
            try:
                self.file_system.create_directory(directory)
            except Exception:
                # Le dossier existe déjà, on continue
                pass

    def get_text(self):
        
        try:
            for a in self.soup.find_all("a", href=True, target="_blank"):
                href = a["href"]

                if href.lower() or "document" in str(href) or ".pdf" in str(href) or ".docx" in str(href) or "doc_num.php" in str(href):
                    pdf_url = urljoin(self.url, href)
                    self.pdf_urls.append(pdf_url)
        except Exception as e:
            print("Erreur de l'url: ", e)
        
    
    def download_text(self):
        self.get_text()

        file_list = list_files_in_adls(self.file_system, self.raw_pdf)

        for pdf_url in self.pdf_urls:
            print("Has recovered :", pdf_url)
            try:
                request = requests.get(pdf_url, stream=True)
                request.raise_for_status()
            except Exception as e:
                print("Error web: ", e)
                continue

            # nom du fichier = dernière partie de l'URL
            if ".docx" in pdf_url or "doc_num.php" in pdf_url:

                # essayer de récupérer le vrai nom dans l'en-tête HTTP
                cd = request.headers.get("Content-Disposition")

                if cd:
                    # extraire le nom du fichier depuis l'en-tête
                    filename = re.findall('filename="?(.+)"?', cd)[0]
                else:
                    # fallback si pas d'en-tête → créer un nom propre
                    filename = pdf_url.split("/")[-1]
                    filename = filename.replace("?", "_").replace("=", "_")  # retirer caractères interdits
                    if not filename.lower().endswith(".docx"):
                        filename += ".docx"
            else:
                filename = pdf_url.split("/")[-1]
            
            filename = filename.strip().strip('"').strip("'")
            filename = re.sub(r'[<>:"/\\|?*]', '_',filename)
            filepath = f"{self.raw_pdf}/{filename}"

            try:
                if filename not in file_list:
                    # Télécharger le fichier en mémoire
                    file_data = b""
                    for chunk in request.iter_content(chunk_size=8192):
                        if chunk:
                            file_data += chunk
                    
                    # Upload vers Azure
                    if upload_file_to_adls(self.file_system, filepath, file_data):
                        print(f"→ Sauvegardé dans {ACCOUNT_NAME}/{FILESYSTEM}/{filepath}")
                    else:
                        print(f"Erreur lors de la sauvegarde de {filename}")
                else:
                    print(f"Text: {filename}, already here.")
            except Exception as e:
                print("can't save:", e)

    def pdf_to_txt(self):
        """
        Convertit tous les PDFs d'un dossier en fichiers TXT dans Azure
        
        Args:
            pdf_folder: Dossier contenant les PDFs à convertir (Azure)
            output_folder: Dossier de destination pour les fichiers TXT (Azure)
        """
        
        # Récupérer tous les fichiers PDF depuis Azure
        pdf_files = list_files_in_adls(self.file_system, self.raw_pdf)
        
        if not pdf_files:
            print(f"  Aucun fichier PDF trouvé dans le dossier '{ACCOUNT_NAME}/{FILESYSTEM}/{self.raw_pdf}/'")
            print(f" Placez vos fichiers PDF dans le dossier '{ACCOUNT_NAME}/{FILESYSTEM}/{self.raw_pdf}/' et relancez le script.")
            return
        
        print(f"\n{'='*80}")
        print(f"   CONVERSION PDF → TXT (Azure) ".center(80))
        print(f"{'='*80}\n")
        print(f" Dossier source : {ACCOUNT_NAME}/{FILESYSTEM}/{self.raw_pdf}/")
        print(f" Dossier destination : {ACCOUNT_NAME}/{FILESYSTEM}/{self.output_folder}/")
        print(f" Nombre de PDFs trouvés : {len(pdf_files)}\n")
        print(f"{'='*80}\n")
        
        success_count = 0
        error_count = 0
        
        # liste avec tout les nom des fichier texte de before_clean_data
        text_list = list_files_in_adls(self.file_system, self.output_folder)

        # Convertir chaque PDF
        for i, pdf_name in enumerate(pdf_files, 1):
            pdf_adls_path = f"{self.raw_pdf}/{pdf_name}"
            root, extension = os.path.splitext(pdf_name)
            
            try:
                # Télécharger le fichier depuis Azure
                pdf_bytes = download_file_from_adls(self.file_system, pdf_adls_path)
                if pdf_bytes is None:
                    print(f"[{i}/{len(pdf_files)}]  Erreur lors du téléchargement : {pdf_name}\n")
                    error_count += 1
                    continue
                
                txt_name = pdf_name.replace(extension, '.txt')
                output_path = f"{self.output_folder}/{txt_name}"
                
                if txt_name not in text_list:
                    print(f"[{i}/{len(pdf_files)}]  Conversion de : {pdf_name}")
                    
                    text = ""
                    pages_count = 0
                    
                    # Essayer PDF d'abord
                    try:
                        reader = PdfReader(BytesIO(pdf_bytes))
                        pages_count = len(reader.pages)
                        # Extraire le texte de toutes les pages
                        for page_num, page in enumerate(reader.pages, 1):
                            page_text = page.extract_text() or ""
                            text += page_text
                        
                    except Exception:
                        # Si ce n'est pas un PDF, essayer DOCX
                        try:
                            doc = Document(BytesIO(pdf_bytes))
                            text = ""
                            for paragraph in doc.paragraphs:
                                text += paragraph.text + "\n"
                            pages_count = len(doc.paragraphs)  # Approximation
                        except Exception as e:
                            print(f"    Erreur: format non supporté ou fichier corrompu: {e}")
                            error_count += 1
                            continue
                    
                    # Écrire le texte dans Azure
                    if write_text_to_adls(self.file_system, output_path, text):
                        chars_count = len(text)
                        print(f"    Converti avec succès : {txt_name}")
                        print(f"    Pages : {pages_count} | Caractères : {chars_count:,}\n")
                        success_count += 1
                    else:
                        print(f"    Erreur lors de l'écriture du fichier\n")
                        error_count += 1
                else:
                    print(f"[{i}/{len(pdf_files)}]  {txt_name} existe déjà, ignoré.\n")
                    success_count += 1
                
            except Exception as e:
                print(f"  Erreur lors de la conversion : {str(e)}\n")
                error_count += 1
        
        # Résumé final
        print(f"{'='*80}")
        print(f"   RÉSUMÉ DE LA CONVERSION ".center(80))
        print(f"{'='*80}\n")
        print(f" Conversions réussies : {success_count}")
        print(f" Conversions échouées : {error_count}")
        print(f" Fichiers TXT disponibles dans : {ACCOUNT_NAME}/{FILESYSTEM}/{self.output_folder}/\n")
        print(f"{'='*80}\n")


    def clean_text(self):
        text_path = list_files_in_adls(self.file_system, self.output_folder)

        for text_name in text_path:
            text_directory = f"{self.output_folder}/{text_name}"

            # Lire le texte depuis Azure
            text = read_text_from_adls(self.file_system, text_directory)
            if text is None:
                print(f"Erreur lors de la lecture de {text_name}, ignoré.")
                continue

            # enleve les espace
            text_to_clean = re.sub(r'\s+', ' ', text)
            #enleve les espace devant la ponctuation double
            text_to_clean = re.sub(r'\s([:;?!])', r'\1', text_to_clean)
            # un espace apre ponctuation double
            text_to_clean = re.sub(r'([:;?!])([a-zA-Z0-9])', r'\1 \2', text_to_clean)
            # retire l espace devant la virgule et le point
            text_to_clean = re.sub(r'\s([,\.])', r'\1', text_to_clean)
            # reduit les espace autour des parenthese
            text_to_clean = re.sub(r'\s(\(|\))', r'\1', text_to_clean) # avent ( ou )
            text_to_clean = re.sub(r'(\()\s', r'\1', text_to_clean) # apre ( 
            text_to_clean = re.sub(r'\s(\))', r'\1', text_to_clean) # avant )
            # supprime les espace multiple 
            text_to_clean = re.sub(r'\s{2,}', ' ', text_to_clean)
            # suprime les /
            text_to_clean = text_to_clean.replace("/", "")
            # transmorme tout en minuscule
            text_to_clean = text_to_clean.lower()
            #
            text_to_clean = text_to_clean.strip()

            clean_text_directory = f"{self.final_folder}/{text_name}"

            # Écrire le texte nettoyé dans Azure
            if not write_text_to_adls(self.file_system, clean_text_directory, text_to_clean):
                print(f"Erreur lors de l'écriture de {text_name}")

    def clone_verifie(self):
        # verifie les si il y a des texte en double et les supprime

        text_list = list_files_in_adls(self.file_system, self.final_folder)

        # recupere tout les texte
        for i, text_file in enumerate(text_list):
            try:
                text_directory = f"{self.final_folder}/{text_file}"
                text = read_text_from_adls(self.file_system, text_directory)
                if text is None:
                    continue
            except:
                continue
            
            counter = 0
            for n in text_list:
                try:
                    text_directory_b = f"{self.final_folder}/{n}"
                    text_b = read_text_from_adls(self.file_system, text_directory_b)
                    if text_b is None:
                        continue

                    if text == text_b and counter == 0:
                        counter += 1
                    elif text == text_b and counter > 0:
                        if delete_file_from_adls(self.file_system, text_directory_b):
                            print(f"File {ACCOUNT_NAME}/{FILESYSTEM}/{text_directory_b} removed")
                except:
                    pass

                
        
if __name__ == "__main__":
    # tout les site qu on veut scraper
    # peut avoir les url que vous vouler
    urls = [
        "https://environnement.brussels/pro/gestion-environnementale/gerer-les-dechets/parcours-dechets-professionnels-reduire-trier-et-gerer-vos-dechets-bruxelles"
    ]

    if len(urls) == 0:
        num = int(input("Combien de site a scraper?: "))
        
        for n in range(num):
            url = str(input("Votre url: "))
            urls.append(url)
    for u in urls:
        scrap = TextScrapper(u)
        # telecharge tout les texte
        scrap.download_text()
        scrap.pdf_to_txt()
        scrap.clean_text()

    scrap.clone_verifie()
